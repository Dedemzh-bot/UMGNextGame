#!/usr/bin/env python3
"""Create the neutral NextGame UMG program-document presentation template."""

from __future__ import annotations

import argparse
import os
import zipfile
from datetime import datetime
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor, Twips


BLUE = "2E75B6"
DARK_BLUE = "1F4E79"
LIGHT_BLUE = "E7EFF8"
PALE_BLUE = "F1F4F8"
GRID = "BCCBDA"
TEXT = "202833"
MUTED = "5F6B7A"
WHITE = "FFFFFF"
WIDGET_TREE_HEADERS = ("层级 / Widget", "Class", "Is Variable", "程序用途")
WIDGET_TREE_WIDTHS_MM = (52.0, 34.0, 20.0, 40.0)
WIDGET_TREE_INDENT_TWIPS = 180
WIDGET_TREE_REUSE_ONLY_EMPTY_TEXT = "无自有 WidgetTree 节点（继承结构不重复列出）"
FIXED_TIME = datetime(2000, 1, 1)
FIXED_ZIP_TIME = (2000, 1, 1, 0, 0, 0)


def _set_run_font(run, latin: str = "Arial", east_asia: str = "宋体", size: float | None = None) -> None:
    run.font.name = latin
    if size is not None:
        run.font.size = Pt(size)
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)


def _set_style_font(style, latin: str, east_asia: str, size: float) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    fonts = style._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)


def _get_or_add_style(document: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, style_type)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color: str = TEXT,
    size: float = 9,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    _set_run_font(run, "Arial", "等线", size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)


def _set_fixed_table_geometry(table, widths_mm: tuple[float, ...] | list[float]) -> None:
    if len(widths_mm) != len(table.columns) or any(width <= 0 for width in widths_mm):
        raise ValueError("Fixed table geometry requires one positive width per column.")

    widths_twips = [Mm(width).twips for width in widths_mm]
    table.autofit = False
    table_properties = table._tbl.tblPr

    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.insert(0, table_width)
    table_width.set(qn("w:w"), str(sum(widths_twips)))
    table_width.set(qn("w:type"), "dxa")

    table_layout = table_properties.find(qn("w:tblLayout"))
    if table_layout is None:
        table_layout = OxmlElement("w:tblLayout")
        table_properties.append(table_layout)
    table_layout.set(qn("w:type"), "fixed")

    table_grid = table._tbl.tblGrid
    for grid_column in list(table_grid):
        table_grid.remove(grid_column)
    for width_twips in widths_twips:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width_twips))
        table_grid.append(grid_column)

    for row in table.rows:
        for column_index, cell in enumerate(row.cells):
            cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:w"), str(widths_twips[column_index]))
            cell_width.set(qn("w:type"), "dxa")


def _set_widget_tree_row(row, values: tuple[str, str, str, str], *, depth: int) -> None:
    if depth < 0:
        raise ValueError("WidgetTree depth cannot be negative.")
    for index, value in enumerate(values):
        alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 2 else WD_ALIGN_PARAGRAPH.LEFT
        _set_text(row.cells[index], value, color=MUTED, size=8.8, align=alignment)
    row.cells[0].paragraphs[0].paragraph_format.left_indent = Twips(depth * WIDGET_TREE_INDENT_TWIPS)
    _keep_row_together(row)


def _add_widget_tree_table(document: Document, *, reuse_only_empty: bool = False):
    table = document.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_table_borders(table)
    _set_fixed_table_geometry(table, WIDGET_TREE_WIDTHS_MM)

    for index, header in enumerate(WIDGET_TREE_HEADERS):
        alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 2 else WD_ALIGN_PARAGRAPH.LEFT
        _set_cell_shading(table.rows[0].cells[index], LIGHT_BLUE)
        _set_text(table.rows[0].cells[index], header, bold=True, color=DARK_BLUE, size=9, align=alignment)
    _repeat_header(table.rows[0])
    _keep_row_together(table.rows[0])

    if reuse_only_empty:
        merged_cell = table.rows[1].cells[0].merge(table.rows[1].cells[3])
        _set_text(
            merged_cell,
            WIDGET_TREE_REUSE_ONLY_EMPTY_TEXT,
            color=MUTED,
            size=8.8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _keep_row_together(table.rows[1])
    else:
        _set_widget_tree_row(
            table.rows[1],
            ("[WidgetName]", "[WidgetClass]", "[IsVariable]", "[ProgramPurposeOrBlank]"),
            depth=1,
        )

    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def _add_table(document: Document, headers: list[str], row_values: list[str], widths_mm: list[float] | None = None):
    table = document.add_table(rows=2, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    _set_table_borders(table)
    for index, header in enumerate(headers):
        _set_cell_shading(table.rows[0].cells[index], LIGHT_BLUE)
        _set_text(table.rows[0].cells[index], header, bold=True, color=DARK_BLUE, size=9)
        _set_text(table.rows[1].cells[index], row_values[index], color=MUTED, size=8)
        if widths_mm:
            table.rows[0].cells[index].width = Mm(widths_mm[index])
            table.rows[1].cells[index].width = Mm(widths_mm[index])
    _repeat_header(table.rows[0])
    _keep_row_together(table.rows[0])
    _keep_row_together(table.rows[1])
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def _add_owner_scoped_program_blocks(document: Document, *, owner_label: str) -> None:
    _add_note(
        document,
        f"条件区块：仅在{owner_label}拥有动态集合时保留；只克隆归属于{owner_label}的集合，"
        "每个集合独立一行并明确容器与 EntryClass 关系。",
    )
    collection_label = _add_labeled_paragraph(document, "动态集合：", "[CollectionId]", after=4)
    collection_label.paragraph_format.keep_with_next = True
    _add_table(
        document,
        ["集合标识", "容器控件", "EntryClass", "程序用途"],
        ["[CollectionId]", "[CollectionWidget]", "[EntryWidgetClass]", "[CollectionPurpose]"],
        [24, 30, 66, 26],
    )

    _add_note(
        document,
        f"条件区块：仅在{owner_label}拥有已接受且通过实际变量门控的状态模型时保留；"
        f"只克隆归属于{owner_label}的状态项。",
    )
    state_label = _add_labeled_paragraph(document, "状态说明：", "[StateModelId]", after=4)
    state_label.paragraph_format.keep_with_next = True
    _add_labeled_paragraph(document, "实现策略：", "[ImplementationStrategy]", after=4)
    _add_labeled_paragraph(document, "控制意图：", "[StateControlIntent]", after=4)
    _add_table(
        document,
        ["状态轴", "状态标识", "状态名", "默认", "目标分支或结果"],
        ["[StateAxis]", "[StateId]", "[StateName]", "[DefaultState]", "[StateTarget]"],
        [27, 34, 24, 16, 45],
    )

    _add_note(
        document,
        f"条件区块：仅在{owner_label}有跨资产定位依赖时保留；"
        "支持依赖随使用它的资产记录，不再生成文档级独立章节。",
    )
    _add_labeled_paragraph(document, "支持依赖：", "[SupportDependencyId]", after=4)
    _add_table(
        document,
        ["Parent Class", "依赖定位节点"],
        ["[DependencyParentClass]", "[DependencyLocatorNodes]"],
        [42, 104],
    )


def _add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def _add_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Template Note")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(text)


def _add_labeled_paragraph(document: Document, label: str, value: str, *, after: float = 8):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(after)
    label_run = paragraph.add_run(label)
    _set_run_font(label_run, "Arial", "黑体", 10.5)
    label_run.bold = True
    value_run = paragraph.add_run(value)
    _set_run_font(value_run, "Arial", "宋体", 10.5)
    return paragraph


def _add_callout(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    _set_cell_shading(cell, PALE_BLUE)
    _set_text(cell, text, color=TEXT, size=10)
    _keep_row_together(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _set_table_borders(table, color: str = GRID, size: int = 5) -> None:
    table_properties = table._tbl.tblPr
    borders = table_properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def _add_placeholder_block(document: Document, text: str, *, monospace: bool = False) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    _set_cell_shading(cell, PALE_BLUE)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_together = True
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        _set_run_font(
            run,
            "Consolas" if monospace else "Arial",
            "等线" if monospace else "宋体",
            8 if monospace else 9,
        )
        run.font.color.rgb = RGBColor.from_string(MUTED)
    _set_cell_margins(cell, top=140, start=160, bottom=140, end=160)
    _keep_row_together(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_page_number(paragraph, *, prefix: bool = True) -> None:
    if prefix:
        paragraph.add_run("NextGame UI · ")
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char_begin, instr_text, fld_char_end))


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.55)
    section.footer_distance = Inches(0.55)

    styles = document.styles
    normal = styles["Normal"]
    _set_style_font(normal, "Arial", "宋体", 10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    _set_style_font(title, "Arial", "黑体", 22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.space_after = Pt(10)

    heading1 = styles["Heading 1"]
    _set_style_font(heading1, "Arial", "黑体", 22)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 0, 0)
    heading1.paragraph_format.space_before = Pt(17)
    heading1.paragraph_format.space_after = Pt(10)
    heading1.paragraph_format.keep_with_next = True

    heading2 = styles["Heading 2"]
    _set_style_font(heading2, "Arial", "黑体", 16)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 0, 0)
    heading2.paragraph_format.space_before = Pt(13)
    heading2.paragraph_format.space_after = Pt(8)
    heading2.paragraph_format.keep_with_next = True

    tool_section = _get_or_add_style(document, "Tool Section")
    tool_section.base_style = heading2
    _set_style_font(tool_section, "Arial", "微软雅黑", 17)
    tool_section.font.bold = True
    tool_section.font.color.rgb = RGBColor.from_string(BLUE)
    tool_section.paragraph_format.space_before = Pt(14)
    tool_section.paragraph_format.space_after = Pt(8)
    tool_section.paragraph_format.keep_with_next = True

    tool_asset = _get_or_add_style(document, "Tool Asset")
    tool_asset.base_style = heading2
    _set_style_font(tool_asset, "Arial", "微软雅黑", 13)
    tool_asset.font.bold = True
    tool_asset.font.color.rgb = RGBColor.from_string(BLUE)
    tool_asset.paragraph_format.space_before = Pt(10)
    tool_asset.paragraph_format.space_after = Pt(6)
    tool_asset.paragraph_format.keep_with_next = True

    caption = _get_or_add_style(document, "Tree Caption")
    caption.base_style = normal
    _set_style_font(caption, "Arial", "等线", 8)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(5)

    if "Template Note" not in styles:
        note = styles.add_style("Template Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = styles["Template Note"]
    _set_style_font(note, "Arial", "等线", 8)
    note.font.italic = True
    note.font.color.rgb = RGBColor.from_string(MUTED)
    note.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    header_run = header.add_run("UGame [SystemFolder] · 程序接入说明")
    _set_run_font(header_run, "Arial", "微软雅黑", 8)
    header_run.bold = True
    header_run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.style = styles["Normal"]
    footer_run = footer.add_run("[SystemFolder] · ")
    _set_run_font(footer_run, "Arial", "等线", 8)
    footer_run.font.color.rgb = RGBColor.from_string(MUTED)
    _add_page_number(footer, prefix=False)
    for run in footer.runs:
        _set_run_font(run, "Arial", "等线", 8)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def _normalize_docx(path: Path) -> None:
    normalized = path.with_suffix(".normalized.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        normalized, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as target:
        for source_info in source.infolist():
            if source_info.filename.startswith("customXml/"):
                continue
            data = source.read(source_info.filename)
            if source_info.filename == "[Content_Types].xml":
                root = ET.fromstring(data)
                for child in list(root):
                    if child.attrib.get("PartName", "").startswith("/customXml/"):
                        root.remove(child)
                ET.register_namespace("", "http://schemas.openxmlformats.org/package/2006/content-types")
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif source_info.filename == "word/_rels/document.xml.rels":
                root = ET.fromstring(data)
                for child in list(root):
                    if child.attrib.get("Type", "").endswith("/customXml"):
                        root.remove(child)
                ET.register_namespace("", "http://schemas.openxmlformats.org/package/2006/relationships")
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            info = zipfile.ZipInfo(source_info.filename, FIXED_ZIP_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = source_info.create_system
            info.external_attr = source_info.external_attr
            info.internal_attr = source_info.internal_attr
            target.writestr(info, data)
    os.replace(normalized, path)


def build_template(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_document(document)

    properties = document.core_properties
    properties.title = "NextGame UMG 界面说明中性模板"
    properties.subject = "NextGame UMG 程序接入文档表现层模板"
    properties.author = "NextGame UI"
    properties.last_modified_by = "NextGame UI"
    properties.created = FIXED_TIME
    properties.modified = FIXED_TIME
    properties.revision = 1
    properties.identifier = "nextgame-ui-program-document-template"
    properties.keywords = "NextGame, UMG, program handoff, template"

    title = document.add_paragraph("[YYYYMMDD]_UGame[System]界面说明", style="Heading 1")
    title.paragraph_format.space_after = Pt(14)

    _add_labeled_paragraph(document, "系统目录：", "[SystemFolder]")
    _add_labeled_paragraph(document, "交付日期：", "[DeliveryDate]")
    _add_labeled_paragraph(document, "目标资产：", "[TargetAssetCount] 个")
    _add_labeled_paragraph(document, "交接标识：", "[HandoffId]", after=10)
    _add_callout(document, "程序范围：[ScopeStatement]")

    document.add_paragraph("1. 资产详细说明", style="Tool Section")
    _add_note(document, "将本资产块按目标资产数量完整克隆；主界面优先，随后按依赖顺序排列子控件与条目。")
    intro = document.add_paragraph(
        "以下内容按资产给出实际 WidgetTree 与程序用途合并表，再说明程序接入关系。"
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    intro.paragraph_format.space_after = Pt(8)
    document.add_paragraph("[AssetName][ChineseFunctionalName]", style="Heading 2")
    _add_labeled_paragraph(document, "资产路径：", "[TargetAssetPath]", after=4)
    _add_labeled_paragraph(document, "Parent Class：", "[ParentClass]", after=4)
    _add_labeled_paragraph(document, "功能说明：", "[FunctionalSummary]", after=6)
    _add_note(
        document,
        "按结构化行逐节点克隆；首列用 depth × indentTwipsPerDepth 设置 Word 左缩进，第三列只写小写 true/false，"
        "第四列写同资产程序用途或留空；reuse-only 空树改用固定跨四列说明行。",
    )
    caption = document.add_paragraph(
        "[AssetName][ChineseFunctionalName] · Unreal 保存后实际 WidgetTree",
        style="Tree Caption",
    )
    caption.paragraph_format.keep_with_next = True
    _add_widget_tree_table(document)
    _add_labeled_paragraph(document, "程序接入关系：", "[AssetRelationshipSummary]", after=6)
    _add_labeled_paragraph(document, "[ControlGroupName]：", "[ControlGroupExplanation]", after=4)
    _add_owner_scoped_program_blocks(document, owner_label="本资产")

    document.add_page_break()
    document.add_paragraph("2. 其他资产程序说明", style="Tool Section")
    _add_note(
        document,
        "条件模块：仅当动态集合、状态模型或支持依赖的归属资产不在第 1 模块时保留；"
        "按归属资产克隆下面的完整分组，每组只写一次归属资产路径，并省略无内容的子区块。"
        "省略本模块时，后续全局模块必须连续重编号。",
    )
    document.add_paragraph("[OtherAssetName][OtherChineseFunctionalName]", style="Heading 2")
    _add_labeled_paragraph(document, "资产路径：", "[OtherAssetPath]", after=4)
    _add_owner_scoped_program_blocks(document, owner_label="上方其他资产")

    document.add_paragraph("3. 构建偏差", style="Tool Section")
    _add_note(document, "条件区块：没有已接受构建偏差时省略本节。")
    _add_table(document, ["资产", "偏差摘要"], ["[AssetName]", "[AcceptedDeviationSummary]"], [42, 104])

    document.add_paragraph("4. 状态控制待接入项", style="Tool Section")
    _add_note(document, "条件区块：没有明确状态控制缺口时省略本节；项目排除项不得列为缺口。")
    _add_table(document, ["来源", "缺口说明"], ["[SourceIdentifier]", "[StateControlGap]"], [42, 104])

    document.save(output_path)
    _normalize_docx(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", required=True, type=Path, help="Destination .docx path")
    args = parser.parse_args()
    if args.output.suffix.lower() != ".docx":
        parser.error("--output must end with .docx")
    build_template(args.output.resolve())
    print(args.output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
