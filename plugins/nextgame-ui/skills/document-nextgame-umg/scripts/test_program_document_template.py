#!/usr/bin/env python3
"""Regression tests for the packaged neutral program-document template."""

from __future__ import annotations

import importlib.util
import json
import re
import shutil
import struct
import unittest
import uuid
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm
from docx.table import Table
from docx.text.paragraph import Paragraph


SCRIPT_DIR = Path(__file__).resolve().parent
DOCUMENT_SKILL_DIR = SCRIPT_DIR.parent
SKILLS_DIR = DOCUMENT_SKILL_DIR.parent
TEMPLATE_SKILL_DIR = SKILLS_DIR / "artifact-template-nextgame-umg"
REFERENCE_DOCX = TEMPLATE_SKILL_DIR / "assets" / "reference.docx"
PREVIEW_PNG = TEMPLATE_SKILL_DIR / "assets" / "preview.png"
GENERATOR_PATH = SCRIPT_DIR / "create_program_document_template.py"
WIDGET_TREE_HEADERS = ("层级 / Widget", "Class", "Is Variable", "程序用途")
WIDGET_TREE_WIDTHS_MM = (52.0, 34.0, 20.0, 40.0)
WIDGET_TREE_INDENT_TWIPS = 180
WIDGET_TREE_REUSE_ONLY_EMPTY_TEXT = "无自有 WidgetTree 节点（继承结构不重复列出）"
COLLECTION_HEADERS = ("集合标识", "容器控件", "EntryClass", "程序用途")
STATE_MODEL_HEADERS = ("状态轴", "状态标识", "状态名", "默认", "目标分支或结果")
SUPPORT_DEPENDENCY_HEADERS = ("Parent Class", "依赖定位节点")


def document_text(path: Path) -> str:
    document = Document(path)
    values: list[str] = []
    values.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    for section in document.sections:
        values.extend(paragraph.text for paragraph in section.header.paragraphs)
        values.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(values)


def load_generator_module():
    spec = importlib.util.spec_from_file_location("create_program_document_template", GENERATOR_PATH)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot load generator: {GENERATOR_PATH}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def widget_tree_tables(document: Document):
    return [
        table
        for table in document.tables
        if tuple(cell.text for cell in table.rows[0].cells) == WIDGET_TREE_HEADERS
    ]


def tables_with_headers(document: Document, headers: tuple[str, ...]):
    return [
        table
        for table in document.tables
        if tuple(cell.text for cell in table.rows[0].cells) == headers
    ]


def ordered_body_blocks(document: Document) -> list[tuple[str, str]]:
    """Return direct body paragraphs and tables without losing their interleaving."""

    blocks: list[tuple[str, str]] = []
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            blocks.append(("paragraph", Paragraph(child, document).text))
        elif child.tag == qn("w:tbl"):
            table = Table(child, document)
            first_row = " | ".join(cell.text for cell in table.rows[0].cells) if table.rows else ""
            blocks.append(("table", first_row))
    return blocks


class ProgramDocumentTemplateTest(unittest.TestCase):
    def assert_widget_tree_table_contract(self, document: Document) -> None:
        matches = widget_tree_tables(document)
        self.assertEqual(1, len(matches))
        table = matches[0]
        self.assertEqual(4, len(table.columns))
        self.assertGreaterEqual(len(table.rows), 2)

        expected_twips = [Mm(width).twips for width in WIDGET_TREE_WIDTHS_MM]
        table_width = table._tbl.tblPr.find(qn("w:tblW"))
        self.assertIsNotNone(table_width)
        self.assertEqual("dxa", table_width.get(qn("w:type")))
        self.assertEqual(sum(expected_twips), int(table_width.get(qn("w:w"))))
        table_layout = table._tbl.tblPr.find(qn("w:tblLayout"))
        self.assertIsNotNone(table_layout)
        self.assertEqual("fixed", table_layout.get(qn("w:type")))
        self.assertEqual(
            expected_twips,
            [int(column.get(qn("w:w"))) for column in table._tbl.tblGrid.gridCol_lst],
        )
        for row in table.rows:
            self.assertIsNone(row._tr.get_or_add_trPr().find(qn("w:trHeight")))
            self.assertIsNotNone(row._tr.get_or_add_trPr().find(qn("w:cantSplit")))
            self.assertEqual(
                expected_twips,
                [int(cell._tc.get_or_add_tcPr().get_or_add_tcW().get(qn("w:w"))) for cell in row.cells],
            )

        header_property = table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader"))
        self.assertIsNotNone(header_property)
        self.assertIn(header_property.get(qn("w:val")), (None, "1", "true"))
        data_row = table.rows[1]
        self.assertEqual(
            ["[WidgetName]", "[WidgetClass]", "[IsVariable]", "[ProgramPurposeOrBlank]"],
            [cell.text for cell in data_row.cells],
        )
        self.assertEqual(
            WIDGET_TREE_INDENT_TWIPS,
            data_row.cells[0].paragraphs[0].paragraph_format.left_indent.twips,
        )
        self.assertEqual(WD_ALIGN_PARAGRAPH.CENTER, data_row.cells[2].paragraphs[0].alignment)
        self.assertEqual(WD_ALIGN_PARAGRAPH.LEFT, data_row.cells[3].paragraphs[0].alignment)
        widget_name_run = next(run for run in data_row.cells[0].paragraphs[0].runs if run.text)
        self.assertAlmostEqual(8.8, widget_name_run.font.size.pt, delta=0.31)

    def test_template_package_is_complete_and_relative(self) -> None:
        expected = (
            TEMPLATE_SKILL_DIR / "SKILL.md",
            TEMPLATE_SKILL_DIR / "artifact-template.json",
            TEMPLATE_SKILL_DIR / "agents" / "openai.yaml",
            REFERENCE_DOCX,
            PREVIEW_PNG,
        )
        for path in expected:
            self.assertTrue(path.is_file(), path)

        metadata = json.loads((TEMPLATE_SKILL_DIR / "artifact-template.json").read_text(encoding="utf-8"))
        self.assertEqual(1, metadata["schemaVersion"])
        self.assertEqual("document", metadata["kind"])
        self.assertEqual("assets/reference.docx", metadata["reference"])
        self.assertEqual("assets/preview.png", metadata["preview"])
        self.assertFalse(Path(metadata["reference"]).is_absolute())
        self.assertFalse(Path(metadata["preview"]).is_absolute())

    def test_reference_has_expected_geometry_styles_and_section_order(self) -> None:
        document = Document(REFERENCE_DOCX)
        self.assertEqual(1, len(document.sections))
        section = document.sections[0]
        self.assertAlmostEqual(210.0, section.page_width.mm, places=1)
        self.assertAlmostEqual(297.0, section.page_height.mm, places=1)
        self.assertAlmostEqual(1.25, section.left_margin.inches, places=2)
        self.assertAlmostEqual(1.25, section.right_margin.inches, places=2)
        self.assertAlmostEqual(1.0, section.top_margin.inches, places=2)
        self.assertAlmostEqual(1.0, section.bottom_margin.inches, places=2)

        self.assertEqual("000000", str(document.styles["Heading 1"].font.color.rgb))
        self.assertEqual("000000", str(document.styles["Heading 2"].font.color.rgb))
        self.assertEqual("2E75B6", str(document.styles["Tool Section"].font.color.rgb))

        text = document_text(REFERENCE_DOCX)
        section_headings = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.style.name == "Tool Section"
        ]
        self.assertEqual(
            [
                "1. 资产详细说明",
                "2. 其他资产程序说明",
                "3. 构建偏差",
                "4. 状态控制待接入项",
            ],
            section_headings,
        )
        self.assertEqual([], [p.text for p in document.paragraphs if p.style.name == "Tool Asset"])
        self.assertEqual(2, len(tables_with_headers(document, COLLECTION_HEADERS)))
        self.assertEqual(2, len(tables_with_headers(document, STATE_MODEL_HEADERS)))
        self.assertEqual(2, len(tables_with_headers(document, SUPPORT_DEPENDENCY_HEADERS)))
        state_labels = [p for p in document.paragraphs if p.text == "状态说明：[StateModelId]"]
        self.assertEqual(2, len(state_labels))
        self.assertTrue(all(p.style.name == "Normal" for p in state_labels))
        collection_labels = [p for p in document.paragraphs if p.text == "动态集合：[CollectionId]"]
        self.assertEqual(2, len(collection_labels))
        self.assertTrue(all(p.style.name == "Normal" for p in collection_labels))
        self.assertTrue(all(p.paragraph_format.keep_with_next for p in collection_labels))
        self.assertNotIn("程序变量清单", text)
        for forbidden in (
            "1. 目标资产",
            "2. 动态集合与 EntryClass",
            "3. 状态模型",
            "支持依赖定位",
            "只读快照边界",
            "附录 A",
            "[ExactSemanticRelationshipStatement]",
            "内容模式：",
            "范围说明：",
            "只读现状说明：",
            "用户确认：",
            "根控件：",
            "所属资产：",
        ):
            self.assertNotIn(forbidden, text)

        header_cell = next(
            cell
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            if cell.text == "层级 / Widget"
        )
        shading = header_cell._tc.get_or_add_tcPr().find(qn("w:shd"))
        self.assertIsNotNone(shading)
        self.assertEqual("E7EFF8", shading.get(qn("w:fill")))
        self.assert_widget_tree_table_contract(document)

    def test_reference_asset_block_preserves_normative_body_order(self) -> None:
        document = Document(REFERENCE_DOCX)
        blocks = ordered_body_blocks(document)

        scope_index = next(
            index
            for index, (kind, text) in enumerate(blocks)
            if kind == "table" and text == "程序范围：[ScopeStatement]"
        )
        section_index = blocks.index(("paragraph", "1. 资产详细说明"))
        self.assertLess(scope_index, section_index)

        expected = (
            ("paragraph", "[AssetName][ChineseFunctionalName]"),
            ("paragraph", "资产路径：[TargetAssetPath]"),
            ("paragraph", "Parent Class：[ParentClass]"),
            ("paragraph", "功能说明：[FunctionalSummary]"),
            ("paragraph", "[AssetName][ChineseFunctionalName] · Unreal 保存后实际 WidgetTree"),
            ("table", "层级 / Widget | Class | Is Variable | 程序用途"),
            ("paragraph", "程序接入关系：[AssetRelationshipSummary]"),
        )
        positions = [blocks.index(block) for block in expected]
        self.assertEqual(sorted(positions), positions)

    def test_reference_is_neutral_and_contains_no_retained_media(self) -> None:
        text = document_text(REFERENCE_DOCX)
        for forbidden in ("Weapon", "Settlement", "/Game/", "handoff:"):
            self.assertNotIn(forbidden, text)
        for marker in (
            "[SystemFolder]",
            "[ScopeStatement]",
            "[TargetAssetPath]",
            "[AssetName][ChineseFunctionalName]",
            "[ParentClass]",
            "[FunctionalSummary]",
            "[OtherAssetName][OtherChineseFunctionalName]",
            "[OtherAssetPath]",
            "[WidgetName]",
            "[WidgetClass]",
            "[IsVariable]",
            "[ProgramPurposeOrBlank]",
            "[StateModelId]",
            "[SupportDependencyId]",
            "[DependencyParentClass]",
            "[DependencyLocatorNodes]",
        ):
            self.assertIn(marker, text)
        for forbidden in (
            "[AssetRole]",
            "[AssetDisplayName]",
            "[WidgetTreeRoot]",
            "[WidgetTreeNode]",
            "[ExactSemanticRelationshipStatement]",
            "所属资产：",
            "附录 A",
            "只读快照边界",
            "├─",
            "└─",
        ):
            self.assertNotIn(forbidden, text)

        with zipfile.ZipFile(REFERENCE_DOCX) as package:
            names = set(package.namelist())
            relationship_payloads = [
                package.read(name).decode("utf-8", errors="ignore")
                for name in names
                if name.endswith(".rels")
            ]
        self.assertFalse(any(name.startswith("word/media/") for name in names))
        self.assertFalse(any(name.startswith("customXml/") for name in names))
        self.assertNotIn("word/comments.xml", names)
        self.assertFalse(any('TargetMode="External"' in payload for payload in relationship_payloads))

    def test_generator_recreates_the_packaged_reference(self) -> None:
        module = load_generator_module()
        directory = SCRIPT_DIR / f".nextgame-doc-template-{uuid.uuid4().hex}"
        directory.mkdir()
        generated = directory / "reference.docx"
        try:
            module.build_template(generated)
            self.assertEqual(document_text(REFERENCE_DOCX), document_text(generated))
            packaged = Document(REFERENCE_DOCX)
            rebuilt = Document(generated)
            self.assertEqual(packaged.sections[0].page_width, rebuilt.sections[0].page_width)
            self.assertEqual(packaged.sections[0].page_height, rebuilt.sections[0].page_height)
            self.assert_widget_tree_table_contract(rebuilt)
        finally:
            shutil.rmtree(directory, ignore_errors=True)

    def test_generator_builds_native_widget_tree_table(self) -> None:
        module = load_generator_module()
        directory = SCRIPT_DIR / f".nextgame-doc-template-{uuid.uuid4().hex}"
        directory.mkdir()
        generated = directory / "reference.docx"
        try:
            module.build_template(generated)
            document = Document(generated)
            self.assert_widget_tree_table_contract(document)
            text = document_text(generated)
            for forbidden in ("[WidgetTreeRoot]", "[WidgetTreeNode]", "├─", "└─"):
                self.assertNotIn(forbidden, text)
            with zipfile.ZipFile(generated) as package:
                self.assertFalse(any(name.startswith("word/media/") for name in package.namelist()))
        finally:
            shutil.rmtree(directory, ignore_errors=True)

    def test_generator_groups_program_blocks_by_owner_asset(self) -> None:
        module = load_generator_module()
        directory = SCRIPT_DIR / f".nextgame-doc-template-{uuid.uuid4().hex}"
        directory.mkdir()
        generated = directory / "reference.docx"
        try:
            module.build_template(generated)
            document = Document(generated)
            text = document_text(generated)
            section_headings = [
                paragraph.text
                for paragraph in document.paragraphs
                if paragraph.style.name == "Tool Section"
            ]
            self.assertEqual(
                [
                    "1. 资产详细说明",
                    "2. 其他资产程序说明",
                    "3. 构建偏差",
                    "4. 状态控制待接入项",
                ],
                section_headings,
            )
            self.assertEqual([], [p.text for p in document.paragraphs if p.style.name == "Tool Asset"])
            self.assertIn("[AssetName][ChineseFunctionalName]", text)
            self.assertIn("程序范围：[ScopeStatement]", text)
            self.assertIn("资产路径：[TargetAssetPath]", text)
            self.assertIn("Parent Class：[ParentClass]", text)
            self.assertIn("功能说明：[FunctionalSummary]", text)
            self.assertIn("程序接入关系：[AssetRelationshipSummary]", text)
            self.assertIn("[OtherAssetName][OtherChineseFunctionalName]", text)
            self.assertIn("资产路径：[OtherAssetPath]", text)
            self.assertIn("条件区块：仅在本资产拥有动态集合时保留", text)
            self.assertIn("条件区块：仅在本资产拥有已接受且通过实际变量门控的状态模型时保留", text)
            self.assertIn("条件区块：仅在本资产有跨资产定位依赖时保留", text)
            self.assertIn("条件模块：仅当动态集合、状态模型或支持依赖的归属资产不在第 1 模块时保留", text)
            self.assertIn("按归属资产克隆下面的完整分组", text)
            self.assertIn("每组只写一次归属资产路径，并省略无内容的子区块。", text)
            self.assertIn("省略本模块时，后续全局模块必须连续重编号。", text)
            self.assertEqual(2, len(tables_with_headers(document, COLLECTION_HEADERS)))
            self.assertEqual(2, len(tables_with_headers(document, STATE_MODEL_HEADERS)))
            self.assertEqual(2, len(tables_with_headers(document, SUPPORT_DEPENDENCY_HEADERS)))
            self.assertEqual(4, text.count("[CollectionId]"))
            self.assertEqual(2, text.count("[StateId]"))
            self.assertEqual(2, text.count("[StateModelId]"))
            self.assertEqual(2, text.count("[SupportDependencyId]"))
            self.assertEqual(1, text.count("[TargetAssetPath]"))
            self.assertEqual(1, text.count("[OtherAssetPath]"))
            state_labels = [p for p in document.paragraphs if p.text == "状态说明：[StateModelId]"]
            self.assertEqual(2, len(state_labels))
            self.assertTrue(all(p.style.name == "Normal" for p in state_labels))
            collection_labels = [p for p in document.paragraphs if p.text == "动态集合：[CollectionId]"]
            self.assertEqual(2, len(collection_labels))
            self.assertTrue(all(p.style.name == "Normal" for p in collection_labels))
            self.assertTrue(all(p.paragraph_format.keep_with_next for p in collection_labels))
            self.assertNotIn("资产角色", text)
            for forbidden in (
                "1. 目标资产",
                "2. 动态集合与 EntryClass",
                "3. 状态模型",
                "[AssetRole]",
                "[AssetDisplayName]",
                "内容模式：",
                "范围说明：",
                "只读现状说明：",
                "用户确认：",
                "根控件：",
                "近期子控件",
                "所属资产：",
                "支持依赖定位",
                "只读快照边界",
                "附录 A",
                "[ExactSemanticRelationshipStatement]",
            ):
                self.assertNotIn(forbidden, text)
            self.assert_widget_tree_table_contract(document)
        finally:
            shutil.rmtree(directory, ignore_errors=True)

    def test_reuse_only_empty_widget_tree_uses_merged_row(self) -> None:
        module = load_generator_module()
        document = Document()
        table = module._add_widget_tree_table(document, reuse_only_empty=True)
        self.assertEqual(WIDGET_TREE_HEADERS, tuple(cell.text for cell in table.rows[0].cells))
        self.assertEqual(WIDGET_TREE_REUSE_ONLY_EMPTY_TEXT, table.rows[1].cells[0].text)
        grid_span = table.rows[1].cells[0]._tc.get_or_add_tcPr().find(qn("w:gridSpan"))
        self.assertIsNotNone(grid_span)
        self.assertEqual("4", grid_span.get(qn("w:val")))
        self.assertIsNotNone(table.rows[1]._tr.get_or_add_trPr().find(qn("w:cantSplit")))
        self.assertIsNone(table.rows[1]._tr.get_or_add_trPr().find(qn("w:trHeight")))

    def test_preview_is_a_rendered_portrait_page(self) -> None:
        payload = PREVIEW_PNG.read_bytes()
        self.assertTrue(payload.startswith(b"\x89PNG\r\n\x1a\n"))
        width, height = struct.unpack(">II", payload[16:24])
        self.assertGreaterEqual(width, 900)
        self.assertGreaterEqual(height, 1200)
        self.assertGreater(height, width)

    def test_document_skill_uses_template_without_machine_paths(self) -> None:
        document_skill = (DOCUMENT_SKILL_DIR / "SKILL.md").read_text(encoding="utf-8")
        template_contract = (DOCUMENT_SKILL_DIR / "references" / "program-document-template.md").read_text(
            encoding="utf-8"
        )
        template_skill = (TEMPLATE_SKILL_DIR / "SKILL.md").read_text(encoding="utf-8")
        template_agent = (TEMPLATE_SKILL_DIR / "agents" / "openai.yaml").read_text(encoding="utf-8")

        for required in (
            "canonical presentation template aligned with `program-document-content.json` 0.4",
            "does not supply business facts or authorize this stage",
            "references/program-document-template.md",
            "../artifact-template-nextgame-umg/artifact-template.json",
            "program-document-content.json` 0.4",
            "--document-content <program-document-content.json>",
            "depth × widgetTreeTables.indentTwipsPerDepth",
            "lowercase `true` or `false`",
            "never emit `w:trHeight`",
            "UIProgramHandoff",
            "Keep one concise `程序范围` callout",
            "Do not create a standalone target-asset section or target-asset table",
            "clone complete asset-detail blocks",
            "uw_fight_use道具使用图标",
            "do not output a `根控件` field",
            "title → `资产路径` → `Parent Class` → `功能说明` → merged WidgetTree/program-purpose table → `程序接入关系`",
            "widgetTreeTables.assets[].parentClassPath",
            "Route dynamic collections, state models, and support-dependency notes by owning or hosting asset",
            "collection.teammate.entries",
            "state-model.teammate-life",
            "ordinary body text or an inline label",
            "Do not repeat `所属资产`",
            "conditional `其他资产程序说明` module",
            "omit the module when empty",
            "never generate document-level collection, state, or `支持依赖定位` sections",
            "embedded collection/state/support-dependency rows",
            "Do not generate a developer-facing `只读快照边界` section or a semantic-trace appendix",
            "requiredSemanticRelationshipStatements",
            "machine-side contracts and verification artifacts only",
            "native Word table",
            "`层级 / Widget`, `Class`, `Is Variable`, and `程序用途`",
            "Never generate a separate document-level program-variable section",
            "无自有 WidgetTree 节点（继承结构不重复列出）",
        ):
            self.assertIn(required, document_skill)
        self.assertNotIn("pass the resulting diagram files", document_skill)
        self.assertNotIn("clone target rows", document_skill)

        for required in (
            "Normative status and authority",
            "canonical structure and presentation contract",
            "Authority order is",
            "A downstream artifact may specialize verified content but may not redefine upstream semantics",
            "template Skill, UI metadata, neutral reference, preview, generator, strict validator, regression tests",
            "A4 portrait",
            "`程序范围` callout",
            "Do not print `内容模式`",
            "Do not create a standalone target-assets section or target-asset table",
            "asset details",
            "Program variables remain inside each asset-detail table",
            "uw_fight_use道具使用图标",
            "do not emit a `根控件` field",
            "the exact-basename-plus-Chinese title; `资产路径`; `Parent Class`; `功能说明`",
            "widgetTreeTables.assets[].parentClassPath",
            "Dynamic collections, state models, and support-dependency notes are routed into asset blocks",
            "collection.teammate.entries",
            "never form standalone document-level sections",
            "state-model.teammate-life",
            "ordinary body text or an inline label",
            "Do not repeat `所属资产`",
            "conditional `其他资产程序说明` module",
            "Omit the entire module when there are no unmatched items",
            "dynamic collections",
            "state models",
            "support-dependency",
            "accepted build deviations",
            "handoff gaps",
            "Do not add a developer-facing `只读快照边界` section or a semantic-trace appendix",
            "requiredSemanticRelationshipStatements",
            "machine contracts and verification artifacts",
            "native Word table",
            "`52 mm`, `34 mm`, `20 mm`, and `40 mm`",
            "depth × widgetTreeTables.indentTwipsPerDepth",
            "lowercase `true` or `false`",
            "do not emit any `w:trHeight`",
            "无自有 WidgetTree 节点（继承结构不重复列出）",
        ):
            self.assertIn(required, template_contract)

        for required in (
            "Positioning and authority",
            "canonical presentation template for NextGame production UMG programmer handoff documents",
            "not a requirement-analysis template, design-review template, prototype/legacy output template",
            "validated `UIProgramHandoff` and `program-document-content.json` are business-content authority",
            "Never copy a concrete system document back into the neutral template",
            "UI metadata, retained reference, preview, generator, strict validators, regression tests",
            "retain one concise `程序范围` callout",
            "Do not create a standalone target-asset section or target-asset table",
            "uw_fight_use道具使用图标",
            "do not output a `根控件` field",
            "title → `资产路径` → `Parent Class` → `功能说明` → merged WidgetTree/program-purpose table → `程序接入关系`",
            "Populate `Parent Class` from the verified 0.4 content contract",
            "Do not print `内容模式`",
            "Route dynamic collections, state models, and support-dependency notes by their owning or hosting asset",
            "collection.teammate.entries",
            "state-model.teammate-life",
            "ordinary body text or an inline label",
            "Do not repeat `所属资产`",
            "conditional `其他资产程序说明` module",
            "Omit the whole module when there are no unmatched items",
            "Do not emit a developer-facing `只读快照边界` section, Appendix A, or a semantic-trace appendix",
            "requiredSemanticRelationshipStatements",
            "machine-side validation data only",
            "native Word table",
            "`层级 / Widget`, `Class`, `Is Variable`, and `程序用途`",
            "`52 / 34 / 20 / 40 mm`",
            "depth × widgetTreeTables.indentTwipsPerDepth",
            "lowercase `true` or `false`",
            "do not emit any `w:trHeight`",
            "无自有 WidgetTree 节点（继承结构不重复列出）",
        ):
            self.assertIn(required, template_skill)

        self.assertIn('short_description: "Create asset-scoped NextGame UMG 0.4 handoffs"', template_agent)
        self.assertIn("native four-column WidgetTree tables", template_agent)

        machine_path = re.compile(r"(?i)(?:\b[A-Z]:[\\/]|\\\\[^\s]+\\|/Users/|/home/|\.codex[\\/]plugins[\\/]cache)")
        for name, content in (
            ("document skill", document_skill),
            ("template contract", template_contract),
            ("template skill", template_skill),
            ("template agent metadata", template_agent),
        ):
            self.assertIsNone(machine_path.search(content), name)


if __name__ == "__main__":
    unittest.main()
