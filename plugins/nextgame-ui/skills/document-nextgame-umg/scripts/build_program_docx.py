#!/usr/bin/env python3
"""Deterministically build a NextGame UMG programmer handoff DOCX.

The builder revalidates the current post-build source chain, accepts business
facts only from UIProgramHandoff 0.3 and its exact program-document-content 0.4
projection, and uses the retained neutral DOCX only for presentation parts.
"""

from __future__ import annotations

import argparse
import importlib.util
import json
import os
import shutil
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Sequence
from xml.etree import ElementTree

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from _document_contract_common import (
    HANDOFF_SCHEMA,
    PROGRAM_DOCUMENT_CONTENT_SCHEMA,
    canonical_sha256,
    load_json,
    sha256_file,
    validate_schema_instance,
)
from prepare_program_document_contract import build_document_content_contract
from validate_program_handoff import validate_program_handoff


EXPECTED_HANDOFF_VERSION = "0.3"
EXPECTED_CONTENT_VERSION = "0.4"
EXPECTED_ROOT_ENVIRONMENT_VARIABLE = "NEXTGAME_UI_PROGRAM_DOCS_ROOT"
EXPECTED_TREE_FORMAT = "word-native-four-column-asset-detail-table-v2"
EXPECTED_TREE_HEADERS = ("层级 / Widget", "Class", "Is Variable", "程序用途")
EXPECTED_TREE_WIDTHS_MM = (52.0, 34.0, 20.0, 40.0)
EXPECTED_TREE_INDENT_TWIPS = 180
REUSE_ONLY_EMPTY_TEXT = "无自有 WidgetTree 节点（继承结构不重复列出）"
FIXED_CORE_TIME = datetime(2000, 1, 1)
PRESERVE_PRESENTATION_PARTS = (
    "docProps/app.xml",
    "docProps/thumbnail.jpeg",
    "word/styles.xml",
    "word/stylesWithEffects.xml",
    "word/numbering.xml",
    "word/theme/theme1.xml",
    "word/fontTable.xml",
    "word/settings.xml",
    "word/webSettings.xml",
)


@dataclass(frozen=True)
class ProgramDocxBuildResult:
    output: Path
    changed: bool


def load_template_helpers(path: Path):
    spec = importlib.util.spec_from_file_location("nextgame_program_document_template", path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Unable to load template helpers: {path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def resolve_retained_reference() -> tuple[Path, Path]:
    document_skill_root = Path(__file__).resolve().parent.parent
    template_skill_root = document_skill_root.parent / "artifact-template-nextgame-umg"
    manifest_path = template_skill_root / "artifact-template.json"
    manifest = load_json(manifest_path)
    if set(manifest) != {"schemaVersion", "kind", "reference", "preview"}:
        raise ValueError("Artifact-template manifest does not match its closed shape.")
    if manifest["schemaVersion"] != 1 or manifest["kind"] != "document":
        raise ValueError("Artifact-template manifest identity is unsupported.")
    raw_reference = manifest["reference"]
    if not isinstance(raw_reference, str) or not raw_reference:
        raise ValueError("Artifact-template reference must be a non-empty relative path.")
    reference = (template_skill_root / raw_reference).resolve()
    try:
        reference.relative_to(template_skill_root.resolve())
    except ValueError as error:
        raise ValueError("Artifact-template reference escapes its skill directory.") from error
    if not reference.is_file() or reference.suffix.lower() != ".docx":
        raise FileNotFoundError(f"Retained DOCX reference is unavailable: {reference}")
    return reference, document_skill_root / "scripts" / "create_program_document_template.py"


def validate_current_inputs(
    *,
    handoff: dict[str, Any],
    handoff_path: Path,
    content: dict[str, Any],
    requirement: dict[str, Any],
    requirement_path: Path,
    bundle: dict[str, Any],
    bundle_path: Path,
    readback: dict[str, Any],
    readback_path: Path,
    build_acceptance: dict[str, Any],
    build_acceptance_path: Path,
) -> None:
    handoff_report = validate_program_handoff(
        handoff,
        load_json(HANDOFF_SCHEMA),
        handoff_path=handoff_path,
        requirement=requirement,
        requirement_path=requirement_path,
        bundle=bundle,
        bundle_path=bundle_path,
        readback=readback,
        readback_path=readback_path,
        build_acceptance=build_acceptance,
        build_acceptance_path=build_acceptance_path,
    )
    if not handoff_report["valid"]:
        codes = sorted({problem.get("code", "invalid") for problem in handoff_report["errors"]})
        raise ValueError(f"UIProgramHandoff is not the exact current accepted projection: {codes}")
    content_errors = validate_schema_instance(content, load_json(PROGRAM_DOCUMENT_CONTENT_SCHEMA))
    if content_errors or content.get("version") != EXPECTED_CONTENT_VERSION:
        raise ValueError("program-document-content must be a valid 0.4 contract.")
    if content.get("handoff", {}).get("fileName") != handoff_path.name:
        raise ValueError("program-document-content handoff filename binding is stale.")
    if content.get("handoff", {}).get("sha256") != sha256_file(handoff_path):
        raise ValueError("program-document-content handoff hash binding is stale.")
    expected_content = build_document_content_contract(
        handoff,
        handoff_path,
        build_acceptance,
        build_acceptance_path,
        requirement,
        requirement_path,
        bundle,
        bundle_path,
        readback,
        readback_path,
    )
    if canonical_sha256(content) != canonical_sha256(expected_content):
        raise ValueError("program-document-content is not the exact projection of the current accepted sources.")

    tree_contract = content.get("widgetTreeTables", {})
    if tree_contract.get("format") != EXPECTED_TREE_FORMAT:
        raise ValueError("Unexpected WidgetTree table format.")
    if tuple(tree_contract.get("headers", [])) != EXPECTED_TREE_HEADERS:
        raise ValueError("Unexpected WidgetTree table headers.")
    if tree_contract.get("indentTwipsPerDepth") != EXPECTED_TREE_INDENT_TWIPS:
        raise ValueError("Unexpected WidgetTree indentation contract.")
    handoff_ids = [asset.get("assetId") for asset in handoff.get("assets", [])]
    tree_ids = [asset.get("assetId") for asset in tree_contract.get("assets", [])]
    if tree_ids != handoff_ids:
        raise ValueError("Handoff and WidgetTree projection asset order differ.")


def clear_document_body(document: Document) -> None:
    body = document._element.body
    section_properties = body.sectPr
    for child in list(body):
        if child is not section_properties:
            body.remove(child)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def configure_header_footer(document: Document, system_folder: str, helpers) -> None:
    for section in document.sections:
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(f"UGame {system_folder} · 程序接入说明")
        helpers._set_run_font(run, "Arial", "微软雅黑", 8)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(helpers.MUTED)

        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(f"{system_folder} · ")
        helpers._set_run_font(run, "Arial", "等线", 8)
        run.font.color.rgb = RGBColor.from_string(helpers.MUTED)
        helpers._add_page_number(paragraph, prefix=False)
        for footer_run in paragraph.runs:
            helpers._set_run_font(footer_run, "Arial", "等线", 8)
            footer_run.font.color.rgb = RGBColor.from_string(helpers.MUTED)


def add_labeled_paragraph(document: Document, label: str, value: str, helpers, *, after: float = 4, keep: bool = False):
    paragraph = helpers._add_labeled_paragraph(document, label, value, after=after)
    paragraph.paragraph_format.keep_with_next = keep
    return paragraph


def add_compact_table(
    document: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths_mm: Sequence[float],
    helpers,
    *,
    center_columns: tuple[int, ...] = (),
):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    helpers._set_table_borders(table)
    for index, header in enumerate(headers):
        alignment = WD_ALIGN_PARAGRAPH.CENTER if index in center_columns else WD_ALIGN_PARAGRAPH.LEFT
        helpers._set_cell_shading(table.rows[0].cells[index], helpers.LIGHT_BLUE)
        helpers._set_text(table.rows[0].cells[index], header, bold=True, color=helpers.DARK_BLUE, size=9, align=alignment)
    helpers._repeat_header(table.rows[0])
    helpers._keep_row_together(table.rows[0])
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if index in center_columns else WD_ALIGN_PARAGRAPH.LEFT
            helpers._set_text(row.cells[index], str(value), color=helpers.MUTED, size=8.6, align=alignment)
        helpers._keep_row_together(row)
    helpers._set_fixed_table_geometry(table, list(widths_mm))
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_widget_tree_table(document: Document, tree_asset: dict[str, Any], helpers) -> None:
    rows = tree_asset.get("treeRows")
    if not isinstance(rows, list):
        raise ValueError(f"Asset {tree_asset.get('assetId')} has invalid WidgetTree rows.")
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    helpers._set_table_borders(table)
    for index, header in enumerate(EXPECTED_TREE_HEADERS):
        alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 2 else WD_ALIGN_PARAGRAPH.LEFT
        helpers._set_cell_shading(table.rows[0].cells[index], helpers.LIGHT_BLUE)
        helpers._set_text(table.rows[0].cells[index], header, bold=True, color=helpers.DARK_BLUE, size=9, align=alignment)
    helpers._repeat_header(table.rows[0])
    helpers._keep_row_together(table.rows[0])

    if not rows:
        if tree_asset.get("emptyState") != "reuse-only-no-owned-widgets":
            raise ValueError(f"Only a verified reuse-only asset may have no owned WidgetTree rows: {tree_asset.get('assetId')}")
        # Set the four-column grid before merging. Reapplying per-column tcW
        # after gridSpan=4 would write the same merged cell four times and
        # leave only the final 40 mm width instead of the full 146 mm span.
        helpers._set_fixed_table_geometry(table, list(EXPECTED_TREE_WIDTHS_MM))
        row = table.add_row()
        merged = row.cells[0].merge(row.cells[3])
        helpers._set_text(merged, REUSE_ONLY_EMPTY_TEXT, color=helpers.MUTED, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
        helpers._keep_row_together(row)
    else:
        for row_data in rows:
            depth = row_data.get("depth")
            if not isinstance(depth, int) or depth < 0:
                raise ValueError(f"Invalid WidgetTree depth in {tree_asset.get('assetId')}: {depth!r}")
            row = table.add_row()
            values = (
                str(row_data["widgetName"]),
                str(row_data["className"]),
                "true" if row_data.get("isVariable") is True else "false",
                str(row_data.get("programPurpose", "")),
            )
            helpers._set_widget_tree_row(row, values, depth=depth)
        helpers._set_fixed_table_geometry(table, list(EXPECTED_TREE_WIDTHS_MM))
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def asset_function_name(asset: dict[str, Any], tree_asset: dict[str, Any]) -> str:
    basename = str(tree_asset.get("assetPath", "")).rsplit("/", 1)[-1]
    parent_class = str(tree_asset.get("parentClassPath", ""))
    if basename.startswith("umg_"):
        return "系统主界面"
    if parent_class.endswith(".ListViewItem"):
        return "集合条目控件"
    if basename.startswith("uw_"):
        return "功能子控件"
    return "界面控件"


def asset_functional_summary(asset: dict[str, Any]) -> str:
    parts: list[str] = []
    variables = asset.get("programVariables", [])
    collections = asset.get("collections", [])
    states = asset.get("states", [])
    if variables:
        parts.append(f"包含 {len(variables)} 个程序变量，精确定位与用途见 WidgetTree 表")
    if collections:
        parts.append(f"承载 {len(collections)} 个由程序填充的动态集合")
    if states:
        parts.append(f"承载 {len(states)} 个高层状态模型")
    if not parts:
        return "该资产当前没有程序读写项，仅保留已验证的 WidgetTree 定位信息。"
    return "；".join(parts) + "。"


def asset_relationship_summary(asset: dict[str, Any]) -> str:
    relationships: list[str] = []
    variable_count = len(asset.get("programVariables", []))
    collection_count = len(asset.get("collections", []))
    state_count = len(asset.get("states", []))
    if variable_count:
        relationships.append(f"{variable_count} 个程序变量的精确定位与用途仅见上表")
    if collection_count:
        relationships.append(f"{collection_count} 个动态集合的容器、EntryClass 与用途见下方集合说明")
    if state_count:
        relationships.append(f"{state_count} 个状态模型的控制意图与目标分支见下方状态说明")
    return "；".join(relationships) + "。" if relationships else "当前无程序读写或状态控制关系。"


def state_target(state: dict[str, Any], implementation_strategy: str) -> str:
    outcomes = state.get("runtimeVisibilityOutcomes", [])
    if outcomes:
        return "；".join(f"{outcome['widgetName']}={outcome['visibility']}" for outcome in outcomes)
    if implementation_strategy == "exclusive-panel-branches":
        branches = state.get("actualSavedVisibilityBindings", [])
        return "；".join(f"目标分支：{branch['widgetName']}" for branch in branches)
    return ""


def add_collection(document: Document, collection: dict[str, Any], helpers) -> None:
    add_labeled_paragraph(document, "动态集合：", collection["id"], helpers, after=4, keep=True)
    add_compact_table(
        document,
        ("集合标识", "容器控件", "EntryClass", "程序用途"),
        ((collection["id"], collection["widgetName"], collection["entryWidgetClass"], collection["purpose"]),),
        (24, 30, 66, 26),
        helpers,
    )


def add_state_model(document: Document, model: dict[str, Any], helpers) -> None:
    add_labeled_paragraph(document, "状态说明：", model["id"], helpers, after=4, keep=True)
    add_labeled_paragraph(document, "实现策略：", model["implementationStrategy"], helpers, after=4, keep=True)
    controls = [
        f"{control['id']}（{control['kind']}）：{control['description']}；目标状态：{'、'.join(control.get('targetStateIds', []))}"
        for control in model.get("controlInputs", [])
    ]
    add_labeled_paragraph(document, "控制意图：", "；".join(controls) if controls else "未声明控制输入", helpers, after=4, keep=True)
    rows: list[tuple[str, str, str, str, str]] = []
    for axis in model.get("axes", []):
        for state in axis.get("states", []):
            rows.append(
                (
                    axis["id"],
                    state["id"],
                    state["name"],
                    "true" if state.get("isDefault") is True else "false",
                    state_target(state, model["implementationStrategy"]),
                )
            )
    add_compact_table(
        document,
        ("状态轴", "状态标识", "状态名", "默认", "目标分支或结果"),
        rows,
        (27, 34, 24, 16, 45),
        helpers,
        center_columns=(3,),
    )


def add_asset_block(document: Document, asset: dict[str, Any], tree_asset: dict[str, Any], helpers) -> None:
    asset_path = tree_asset["assetPath"]
    basename = asset_path.rsplit("/", 1)[-1]
    function_name = asset_function_name(asset, tree_asset)
    heading = document.add_paragraph(f"{basename}{function_name}", style="Heading 2")
    heading.paragraph_format.keep_with_next = True
    add_labeled_paragraph(document, "资产路径：", asset_path, helpers, after=4, keep=True)
    add_labeled_paragraph(document, "Parent Class：", tree_asset["parentClassPath"], helpers, after=4, keep=True)
    add_labeled_paragraph(document, "功能说明：", asset_functional_summary(asset), helpers, after=6, keep=True)
    caption = document.add_paragraph(f"{basename}{function_name} · WidgetTree", style="Tree Caption")
    caption.paragraph_format.keep_with_next = True
    add_widget_tree_table(document, tree_asset, helpers)
    add_labeled_paragraph(document, "程序接入关系：", asset_relationship_summary(asset), helpers, after=6)
    for collection in asset.get("collections", []):
        add_collection(document, collection, helpers)
    for model in asset.get("states", []):
        add_state_model(document, model, helpers)


def restore_presentation_parts(reference: Path, authored: Path) -> None:
    with zipfile.ZipFile(reference, "r") as reference_zip, zipfile.ZipFile(authored, "r") as authored_zip:
        reference_names = set(reference_zip.namelist())
        authored_names = set(authored_zip.namelist())
        missing = [name for name in PRESERVE_PRESENTATION_PARTS if name not in reference_names or name not in authored_names]
        if missing:
            raise ValueError(f"Presentation package parts are missing: {missing}")
        source_parts = {name: authored_zip.read(name) for name in authored_zip.namelist()}
        for name in PRESERVE_PRESENTATION_PARTS:
            source_parts[name] = reference_zip.read(name)

    restored = authored.with_suffix(".restored.docx")
    with zipfile.ZipFile(restored, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as target:
        for name in sorted(source_parts):
            target.writestr(name, source_parts[name])
    os.replace(restored, authored)


def visible_docx_text(path: Path) -> str:
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    parts: list[str] = []
    with zipfile.ZipFile(path, "r") as archive:
        for name in sorted(archive.namelist()):
            if name == "word/document.xml" or name.startswith("word/header") or name.startswith("word/footer"):
                root = ElementTree.fromstring(archive.read(name))
                parts.extend(node.text or "" for node in root.iter(f"{namespace}t"))
    return "\n".join(parts)


def self_check(path: Path, content: dict[str, Any], expected_title: str) -> None:
    text = visible_docx_text(path)
    if expected_title not in text:
        raise RuntimeError("Output title is missing.")
    forbidden = ("[SystemFolder]", "[DeliveryDate]", "[AssetName]", "只读快照边界", "Appendix A", "附录 A")
    present = [token for token in forbidden if token in text]
    if present:
        raise RuntimeError(f"Forbidden template or machine-only text remains: {present}")
    for field, identifiers in content.get("requiredIdentifiers", {}).items():
        if isinstance(identifiers, list):
            missing = [identifier for identifier in identifiers if isinstance(identifier, str) and identifier not in text]
            if missing:
                raise RuntimeError(f"Missing required identifiers for {field}: {missing}")
    leaked_statements = [
        statement
        for statement in content.get("requiredSemanticRelationshipStatements", [])
        if isinstance(statement, str) and statement in text
    ]
    if leaked_statements:
        raise RuntimeError("Machine-only semantic relationship statements leaked into the DOCX.")


def validate_package_boundary(path: Path, reference: Path, forbidden_paths: Sequence[Path]) -> None:
    with zipfile.ZipFile(reference, "r") as reference_zip, zipfile.ZipFile(path, "r") as archive:
        allowed_parts = {name for name in reference_zip.namelist() if not name.startswith("customXml/")}
        actual_parts = set(archive.namelist())
        if actual_parts != allowed_parts:
            raise RuntimeError(
                f"DOCX package part boundary changed: missing={sorted(allowed_parts - actual_parts)}, "
                f"unexpected={sorted(actual_parts - allowed_parts)}"
            )
        for info in archive.infolist():
            if info.date_time != helpers_fixed_zip_time():
                raise RuntimeError(f"DOCX package part has a volatile timestamp: {info.filename}")
            if info.compress_type != zipfile.ZIP_STORED:
                raise RuntimeError(f"DOCX package part is not stored canonically: {info.filename}")
            if info.extra or info.comment:
                raise RuntimeError(f"DOCX package part has volatile ZIP metadata: {info.filename}")
            if info.filename.endswith(".rels"):
                root = ElementTree.fromstring(archive.read(info.filename))
                if any(node.attrib.get("TargetMode") == "External" for node in root):
                    raise RuntimeError(f"External DOCX relationship is forbidden: {info.filename}")
        xml_text = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in sorted(actual_parts)
            if name.endswith(".xml") or name.endswith(".rels")
        )
        leaked_paths = [str(value.resolve()) for value in forbidden_paths if str(value.resolve()) in xml_text]
        if leaked_paths:
            raise RuntimeError("DOCX package leaked an output or temporary filesystem path.")


def helpers_fixed_zip_time() -> tuple[int, int, int, int, int, int]:
    return (2000, 1, 1, 0, 0, 0)


def resolve_output(handoff: dict[str, Any], output_root_override: Path | None) -> Path:
    output = handoff["output"]
    environment_variable = output.get("rootEnvironmentVariable")
    if environment_variable != EXPECTED_ROOT_ENVIRONMENT_VARIABLE:
        raise ValueError("Unexpected document output-root environment variable.")
    if output_root_override is None:
        raw_root = os.environ.get(environment_variable)
        if not raw_root:
            raise RuntimeError(f"User environment variable {environment_variable} is not configured.")
        output_root = Path(raw_root).expanduser().resolve()
    else:
        output_root = output_root_override.resolve()
    output_root.mkdir(parents=True, exist_ok=True)
    file_name = output.get("fileName")
    if not isinstance(file_name, str) or not file_name or Path(file_name).name != file_name:
        raise ValueError("Output DOCX filename must be one direct-child basename.")
    destination = output_root / file_name
    if destination.parent.resolve() != output_root or destination.suffix.lower() != ".docx":
        raise ValueError("Output DOCX must be a direct child of the configured document root.")
    is_junction = bool(getattr(destination, "is_junction", lambda: False)())
    if destination.is_symlink() or is_junction:
        raise ValueError("Output DOCX may not replace a symbolic link or junction.")
    return destination


def build_program_docx(
    *,
    handoff_path: Path,
    content_path: Path,
    requirement_path: Path,
    bundle_path: Path,
    readback_path: Path,
    build_acceptance_path: Path,
    output_root_override: Path | None = None,
) -> ProgramDocxBuildResult:
    handoff_path = handoff_path.resolve()
    content_path = content_path.resolve()
    requirement_path = requirement_path.resolve()
    bundle_path = bundle_path.resolve()
    readback_path = readback_path.resolve()
    build_acceptance_path = build_acceptance_path.resolve()
    handoff = load_json(handoff_path)
    content = load_json(content_path)
    requirement = load_json(requirement_path)
    bundle = load_json(bundle_path)
    readback = load_json(readback_path)
    build_acceptance = load_json(build_acceptance_path)
    validate_current_inputs(
        handoff=handoff,
        handoff_path=handoff_path,
        content=content,
        requirement=requirement,
        requirement_path=requirement_path,
        bundle=bundle,
        bundle_path=bundle_path,
        readback=readback,
        readback_path=readback_path,
        build_acceptance=build_acceptance,
        build_acceptance_path=build_acceptance_path,
    )
    if handoff.get("version") != EXPECTED_HANDOFF_VERSION:
        raise ValueError(f"Expected UIProgramHandoff {EXPECTED_HANDOFF_VERSION}.")

    reference, template_script = resolve_retained_reference()
    helpers = load_template_helpers(template_script)
    destination = resolve_output(handoff, output_root_override)
    with tempfile.NamedTemporaryFile(prefix=f".{destination.stem}.", suffix=".authoring.docx", dir=destination.parent, delete=False) as handle:
        temporary = Path(handle.name)
    try:
        shutil.copyfile(reference, temporary)
        document = Document(str(temporary))
        clear_document_body(document)
        system_folder = handoff["target"]["systemFolder"]
        configure_header_footer(document, system_folder, helpers)
        properties = document.core_properties
        properties.title = destination.stem
        properties.subject = "NextGame UMG 程序接入说明"
        properties.author = "NextGame UI"
        properties.last_modified_by = "NextGame UI"
        properties.created = FIXED_CORE_TIME
        properties.modified = FIXED_CORE_TIME
        properties.revision = 1
        properties.identifier = handoff["handoffId"]
        properties.keywords = "NextGame, UMG, program handoff"

        date_token = destination.name[:8]
        delivery_date = f"{date_token[:4]}-{date_token[4:6]}-{date_token[6:8]}"
        title = document.add_paragraph(destination.stem, style="Heading 1")
        title.paragraph_format.space_after = Pt(14)
        add_labeled_paragraph(document, "系统目录：", system_folder, helpers, after=8)
        add_labeled_paragraph(document, "交付日期：", delivery_date, helpers, after=8)
        add_labeled_paragraph(document, "目标资产：", f"{len(handoff['assets'])} 个", helpers, after=8)
        add_labeled_paragraph(document, "交接标识：", handoff["handoffId"], helpers, after=10)
        variable_count = sum(len(asset.get("programVariables", [])) for asset in handoff["assets"])
        collection_count = sum(len(asset.get("collections", [])) for asset in handoff["assets"])
        state_count = sum(len(asset.get("states", [])) for asset in handoff["assets"])
        helpers._add_callout(
            document,
            f"程序范围：本文档记录 {len(handoff['assets'])} 个资产中的 {variable_count} 个程序变量、"
            f"{collection_count} 个动态集合和 {state_count} 个高层状态模型；静态 Designer 配置不作为运行时接口。",
        )

        section = document.add_paragraph("1. 资产详细说明", style="Tool Section")
        section.paragraph_format.keep_with_next = True
        intro = document.add_paragraph("以下按资产记录 WidgetTree 定位信息与程序接入关系。")
        intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
        intro.paragraph_format.space_after = Pt(8)
        assets_by_id = {asset["assetId"]: asset for asset in handoff["assets"]}
        trees_by_id = {asset["assetId"]: asset for asset in content["widgetTreeTables"]["assets"]}
        for asset in handoff["assets"]:
            add_asset_block(document, assets_by_id[asset["assetId"]], trees_by_id[asset["assetId"]], helpers)

        section_number = 2
        deviations = [item for item in handoff.get("deviations", []) if item.get("status") == "accepted"]
        if deviations:
            heading = document.add_paragraph(f"{section_number}. 构建偏差", style="Tool Section")
            heading.paragraph_format.keep_with_next = True
            add_compact_table(
                document,
                ("偏差标识", "影响", "涉及资产", "需求引用"),
                tuple(
                    (
                        item["id"],
                        item["impact"],
                        "、".join(item["affectedAssetIds"]),
                        "、".join(item["affectedRequirementRefs"]),
                    )
                    for item in deviations
                ),
                (34, 18, 42, 52),
                helpers,
            )
            section_number += 1
        gaps = handoff.get("gaps", [])
        if gaps:
            heading = document.add_paragraph(f"{section_number}. 状态控制待接入项", style="Tool Section")
            heading.paragraph_format.keep_with_next = True
            add_compact_table(
                document,
                ("缺口代码", "状态模型", "控制输入", "缺口说明"),
                tuple(
                    (
                        item["code"],
                        item["stateModelId"],
                        item.get("controlInputId", ""),
                        item["description"],
                    )
                    for item in gaps
                ),
                (34, 34, 32, 46),
                helpers,
            )

        document.save(temporary)
        restore_presentation_parts(reference, temporary)
        helpers._normalize_docx(temporary)
        self_check(temporary, content, destination.stem)
        validate_package_boundary(temporary, reference, (destination.parent, temporary.parent))
        changed = not (destination.is_file() and destination.read_bytes() == temporary.read_bytes())
        if not changed:
            temporary.unlink()
        else:
            os.replace(temporary, destination)
    finally:
        temporary.unlink(missing_ok=True)
        temporary.with_suffix(".restored.docx").unlink(missing_ok=True)
        temporary.with_suffix(".normalized.docx").unlink(missing_ok=True)
    return ProgramDocxBuildResult(output=destination, changed=changed)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--handoff", type=Path, required=True)
    parser.add_argument("--document-content", type=Path, required=True)
    parser.add_argument("--requirement", type=Path, required=True)
    parser.add_argument("--bundle", type=Path, required=True)
    parser.add_argument("--readback", type=Path, required=True)
    parser.add_argument("--build-acceptance", type=Path, required=True)
    args = parser.parse_args()
    try:
        build_result = build_program_docx(
            handoff_path=args.handoff,
            content_path=args.document_content,
            requirement_path=args.requirement,
            bundle_path=args.bundle,
            readback_path=args.readback,
            build_acceptance_path=args.build_acceptance,
        )
    except (OSError, ValueError, RuntimeError, KeyError, TypeError, json.JSONDecodeError) as error:
        error_code = "io-error" if isinstance(error, OSError) else "contract-error"
        print(
            json.dumps(
                {"valid": False, "errorCode": error_code, "errorType": type(error).__name__},
                ensure_ascii=False,
            )
        )
        return 1
    output = build_result.output
    output_sha256 = sha256_file(output)
    print(
        json.dumps(
            {
                "valid": True,
                "fileName": output.name,
                "sha256": output_sha256,
                "byteSize": output.stat().st_size,
                "changed": build_result.changed,
            },
            ensure_ascii=False,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
